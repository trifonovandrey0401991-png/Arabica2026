#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для редактирования шаблона РКО .docx и конвертации в PDF
"""

import sys
import json
import os
import re
from docx import Document
from datetime import datetime

def process_rko_template(template_path, output_path, data):
    """
    Обрабатывает шаблон РКО .docx и заменяет плейсхолдеры {field_name} на реальные данные
    
    Args:
        template_path: путь к шаблону .docx
        output_path: путь для сохранения отредактированного .docx
        data: словарь с данными для замены (ключи соответствуют плейсхолдерам)
    """
    try:
        # Загружаем шаблон
        doc = Document(template_path)
        
        # Маппинг плейсхолдеров из шаблона в данные системы
        # Поддерживаем оба формата: старый (field_name) и новый (FIELD_NAME)
        placeholder_mapping = {
            # Новые плейсхолдеры (из заполненного шаблона)
            'FIO': 'fio_receiver',
            'BASIS': 'basis',
            'AMOUNT_WORDS': 'amount_text',
            'RECEIVED_DATE': 'date_text',
            'PASSPORT': 'passport_info',
            'PASSPORT2': 'passport_issuer',
            'IP_SHORT': 'head_name',  # Короткое имя директора с инициалами
            'IP': 'org_name',  # Полное имя директора (будет обработано отдельно для извлечения только имени)
            'INN': 'inn',  # Только ИНН
            'DOC_ID': 'doc_number',
            'DATE': 'doc_date',
            'AMOUNT': 'amount_numeric',
            'SHOP': 'shop_address',  # Адрес без префикса "Фактический адрес:"
            # Старые плейсхолдеры (для обратной совместимости)
            'org_name': 'org_name',
            'org_address': 'org_address',
            'shop_address': 'shop_address',
            'doc_number': 'doc_number',
            'doc_date': 'doc_date',
            'amount_numeric': 'amount_numeric',
            'fio_receiver': 'fio_receiver',
            'basis': 'basis',
            'amount_text': 'amount_text',
            'attachment': 'attachment',
            'head_position': 'head_position',
            'head_name': 'head_name',
            'receiver_amount_text': 'receiver_amount_text',
            'date_text': 'date_text',
            'passport_info': 'passport_info',
            'passport_issuer': 'passport_issuer',
            'cashier_name': 'cashier_name',
        }
        
        # Извлекаем ИНН из org_name, если нужно
        if 'inn' not in data and 'org_name' in data:
            inn_match = re.search(r'ИНН:\s*(\d+)', data.get('org_name', ''))
            if inn_match:
                data['inn'] = inn_match.group(1)
        
        # Извлекаем только имя ИП из org_name для плейсхолдера {IP}
        if 'org_name' in data:
            org_name_full = data['org_name']
            # Убираем "ИНН: ..." для получения только имени
            ip_name = re.sub(r'\s*ИНН:\s*\d+.*', '', org_name_full).strip()
            data['ip_name'] = ip_name
        
        # Функция для замены плейсхолдеров в тексте
        def replace_placeholders(text):
            """Заменяет все плейсхолдеры {field_name} или {{field_name}} на значения из data"""
            # Ищем все плейсхолдеры в формате {field_name} или {{field_name}}
            # Сначала обрабатываем двойные скобки {{...}}, потом одинарные {...}
            pattern_double = r'\{\{(\w+)\}\}'
            pattern_single = r'\{(\w+)\}'
            
            # Заменяем двойные скобки
            matches_double = re.findall(pattern_double, text)
            for placeholder_name in matches_double:
                placeholder = f'{{{{{placeholder_name}}}}}'
                # Получаем ключ данных через маппинг
                data_key = placeholder_mapping.get(placeholder_name, placeholder_name)
                value = data.get(data_key, '')  # Если поле не найдено, заменяем на пустую строку
                
                # Специальная обработка для {IP} - используем только имя без ИНН
                if placeholder_name == 'IP' and 'ip_name' in data:
                    value = data['ip_name']
                
                text = text.replace(placeholder, str(value))
            
            # Заменяем одинарные скобки
            matches_single = re.findall(pattern_single, text)
            for placeholder_name in matches_single:
                placeholder = f'{{{placeholder_name}}}'
                # Получаем ключ данных через маппинг
                data_key = placeholder_mapping.get(placeholder_name, placeholder_name)
                value = data.get(data_key, '')  # Если поле не найдено, заменяем на пустую строку
                
                # Специальная обработка для {IP} - используем только имя без ИНН
                if placeholder_name == 'IP' and 'ip_name' in data:
                    value = data['ip_name']
                
                text = text.replace(placeholder, str(value))
            
            return text
        
        # Функция для замены плейсхолдеров в параграфе с сохранением структуры
        # УЛУЧШЕННАЯ ВЕРСИЯ: НЕ использует para.clear(), что сохраняет выравнивание, отступы и табуляцию
        def replace_in_paragraph(para):
            """
            УЛУЧШЕННАЯ версия: заменяет плейсхолдеры БЕЗ использования para.clear()
            Это сохраняет ВСЮ структуру параграфа: выравнивание, отступы, табуляцию
            """
            # Собираем весь текст параграфа
            full_text = para.text
            
            # Если нет плейсхолдеров, пропускаем
            if not re.search(r'\{(\w+)\}', full_text):
                return
            
            # Заменяем плейсхолдеры в тексте
            new_text = replace_placeholders(full_text)
            
            # Если текст не изменился, ничего не делаем
            if new_text == full_text:
                return
            
            # КЛЮЧЕВОЕ ОТЛИЧИЕ: НЕ используем para.clear()!
            # Вместо этого заменяем текст внутри существующих runs
            
            # Сначала пытаемся заменить в каждом run отдельно
            for run in para.runs:
                if re.search(r'\{(\w+)\}', run.text):
                    run_text = run.text
                    new_run_text = replace_placeholders(run_text)
                    if new_run_text != run_text:
                        run.text = new_run_text
                        # После замены проверяем, что весь параграф заменен
                        if para.text == new_text:
                            return
            
            # Если плейсхолдер разбит на несколько runs или замена неполная,
            # заменяем весь текст первого run, остальные очищаем (но НЕ удаляем параграф!)
            if para.runs:
                # Сохраняем форматирование первого run
                first_run = para.runs[0]
                saved_formatting = {
                    'name': first_run.font.name,
                    'size': first_run.font.size,
                    'bold': first_run.font.bold,
                    'italic': first_run.font.italic,
                }
                
                # Заменяем текст первого run
                first_run.text = new_text
                
                # Очищаем остальные runs (но НЕ удаляем их и НЕ очищаем параграф!)
                for run in para.runs[1:]:
                    run.text = ''
                
                # Восстанавливаем форматирование
                if saved_formatting['name']:
                    first_run.font.name = saved_formatting['name']
                if saved_formatting['size']:
                    first_run.font.size = saved_formatting['size']
                first_run.font.bold = saved_formatting['bold']
                first_run.font.italic = saved_formatting['italic']
        
        # Заменяем плейсхолдеры в параграфах
        for para in doc.paragraphs:
            replace_in_paragraph(para)
        
        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Обрабатываем параграфы внутри ячейки
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)
        
        # Сохраняем отредактированный документ
        doc.save(output_path)
        return True, None
        
    except Exception as e:
        return False, str(e)


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Конвертирует .docx в PDF используя LibreOffice
    
    Args:
        docx_path: путь к .docx файлу
        pdf_path: путь для сохранения PDF
    """
    try:
        import subprocess
        
        # Используем LibreOffice для конвертации
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path) or '.',
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            # LibreOffice создает PDF с тем же именем, но расширением .pdf
            generated_pdf = docx_path.replace('.docx', '.pdf')
            if os.path.exists(generated_pdf):
                # Переименовываем в нужное имя
                if generated_pdf != pdf_path:
                    os.rename(generated_pdf, pdf_path)
                return True, None
            else:
                return False, "PDF файл не был создан"
        else:
            return False, f"Ошибка конвертации: {result.stderr}"
            
    except FileNotFoundError:
        return False, "LibreOffice не установлен. Установите: apt-get install libreoffice"
    except subprocess.TimeoutExpired:
        return False, "Таймаут конвертации"
    except Exception as e:
        return False, str(e)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({'success': False, 'error': 'Недостаточно аргументов'}))
        sys.exit(1)
    
    command = sys.argv[1]
    
    if command == 'process':
        # Обработка шаблона
        if len(sys.argv) < 5:
            print(json.dumps({'success': False, 'error': 'Недостаточно аргументов для process'}))
            sys.exit(1)
        
        template_path = sys.argv[2]
        output_path = sys.argv[3]
        data_json = sys.argv[4]
        
        try:
            data = json.loads(data_json)
            success, error = process_rko_template(template_path, output_path, data)
            
            if success:
                print(json.dumps({'success': True, 'output_path': output_path}))
            else:
                print(json.dumps({'success': False, 'error': error}))
        except Exception as e:
            print(json.dumps({'success': False, 'error': str(e)}))
    
    elif command == 'convert':
        # Конвертация в PDF
        if len(sys.argv) < 4:
            print(json.dumps({'success': False, 'error': 'Недостаточно аргументов для convert'}))
            sys.exit(1)
        
        docx_path = sys.argv[2]
        pdf_path = sys.argv[3]
        
        success, error = convert_docx_to_pdf(docx_path, pdf_path)
        
        if success:
            print(json.dumps({'success': True, 'pdf_path': pdf_path}))
        else:
            print(json.dumps({'success': False, 'error': error}))
    
    else:
        print(json.dumps({'success': False, 'error': f'Неизвестная команда: {command}'}))

