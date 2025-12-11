#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание шаблона РКО на основе HTML структуры из Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(para):
    """Добавляет горизонтальную линию в параграф"""
    para.add_run('_' * 100)

def create_rko_from_html_structure():
    """Создает шаблон РКО на основе HTML структуры"""
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    
    # ========== ЗАГОЛОВОК (справа вверху) ==========
    para_header = doc.add_paragraph()
    para_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_header = para_header.add_run('Унифицированная форма № КО-2\nУтверждена постановлением Госкомстата России от 18.08.98 № 88')
    run_header.font.size = Pt(8)
    run_header.font.name = 'Times New Roman'
    
    # ========== ТАБЛИЦА 1: Организация, ИНН, Адрес, Номер документа, Дата ==========
    table1 = doc.add_table(rows=8, cols=6)
    table1.style = 'Table Grid'
    
    # Настройка ширины колонок (примерно как в HTML)
    # Колонка 1-2: 502pt (376.55pt) - организация
    # Колонка 3: 101pt (75.45pt) - пустая
    # Колонка 4-6: 128pt (96.25pt) - номер и дата
    
    # Строка 1: Организация и ИНН (объединенные ячейки 1-2)
    cell_org = table1.rows[0].cells[0]
    cell_org.merge(table1.rows[0].cells[1])
    para_org = cell_org.paragraphs[0]
    para_org.add_run(' {{IP}} ').font.name = 'Arial'
    para_org.add_run('ИНН: ').font.name = 'Arial'
    para_org.add_run('{{INN}}').font.name = 'Arial'
    
    # Линия подчеркивания
    cell_line1 = table1.rows[1].cells[0]
    cell_line1.merge(table1.rows[1].cells[1])
    cell_line1.text = '_______________________________________________________________________'
    
    # Строка 3: Фактический адрес
    cell_addr = table1.rows[2].cells[0]
    cell_addr.merge(table1.rows[2].cells[1])
    para_addr = cell_addr.paragraphs[0]
    run_addr = para_addr.add_run('Фактический адрес: {{SHOP}}')
    run_addr.font.size = Pt(8)
    run_addr.font.underline = True
    
    # Линия подчеркивания для адреса
    cell_line2 = table1.rows[3].cells[0]
    cell_line2.merge(table1.rows[3].cells[1])
    cell_line2.text = '_______________________________________________________________________'
    
    # Строка 4: Пустая с подписью "(наименование)"
    cell_name1 = table1.rows[4].cells[0]
    cell_name1.merge(table1.rows[4].cells[1])
    para_name1 = cell_name1.paragraphs[0]
    para_name1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name1 = para_name1.add_run('(наименование)')
    run_name1.font.size = Pt(6)
    
    # Строка 5: Пустая с подписью "(наименование организации)"
    cell_name2 = table1.rows[5].cells[0]
    cell_name2.merge(table1.rows[5].cells[1])
    para_name2 = cell_name2.paragraphs[0]
    para_name2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name2 = para_name2.add_run('(наименование организации)')
    run_name2.font.size = Pt(6)
    
    # Строка 6: Пустая
    for col in range(6):
        table1.rows[6].cells[col].text = ''
    
    # Строка 7: Номер документа и Дата
    # Колонка 3: "Номер по порядку" (выравнивание справа)
    cell_num_label = table1.rows[7].cells[2]
    para_num_label = cell_num_label.paragraphs[0]
    para_num_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_num_label = para_num_label.add_run('Номер по порядку')
    run_num_label.font.size = Pt(8)
    
    # Колонка 4-6: Номер документа (объединенные)
    cell_doc_id = table1.rows[7].cells[3]
    cell_doc_id.merge(table1.rows[7].cells[5])
    para_doc_id = cell_doc_id.paragraphs[0]
    para_doc_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_doc_id = para_doc_id.add_run('{{DOC_ID}}')
    run_doc_id.font.size = Pt(8)
    run_doc_id.font.bold = True
    
    # Строка 8: Дата
    # Колонка 3: "По порядку" (выравнивание справа)
    cell_date_label = table1.rows[7].cells[2]
    # Добавляем в ту же ячейку, но на новой строке
    para_date_label = cell_date_label.add_paragraph()
    para_date_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date_label = para_date_label.add_run('По порядку')
    run_date_label.font.size = Pt(8)
    
    # Колонка 4-6: Дата (объединенные)
    cell_date = table1.rows[7].cells[3]
    para_date = cell_date.add_paragraph()
    para_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = para_date.add_run('{{DATE}}')
    run_date.font.size = Pt(8)
    run_date.font.bold = True
    
    # Добавляем пустую строку
    doc.add_paragraph()
    
    # ========== ТАБЛИЦА 2: Сумма ==========
    table2 = doc.add_table(rows=3, cols=9)
    table2.style = 'Table Grid'
    
    # Строка 1: Заголовки
    # Колонка 1-4: "Сумма"
    cell_sum_label = table2.rows[0].cells[0]
    cell_sum_label.merge(table2.rows[0].cells[3])
    para_sum_label = cell_sum_label.paragraphs[0]
    para_sum_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sum_label = para_sum_label.add_run('Сумма')
    run_sum_label.font.size = Pt(8)
    
    # Колонка 5-6: "Рубли, коп. коп."
    cell_rub_label = table2.rows[0].cells[4]
    cell_rub_label.merge(table2.rows[0].cells[5])
    para_rub_label = cell_rub_label.paragraphs[0]
    para_rub_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rub_label = para_rub_label.add_run('Рубли,\nкоп. коп.')
    run_rub_label.font.size = Pt(8)
    
    # Колонка 7-8: "Для отражения операций"
    cell_op_label = table2.rows[0].cells[6]
    cell_op_label.merge(table2.rows[0].cells[7])
    para_op_label = cell_op_label.paragraphs[0]
    para_op_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_op_label = para_op_label.add_run('Для отражения операций')
    run_op_label.font.size = Pt(8)
    
    # Строка 2: Подзаголовки
    # Колонка 1: пустая
    # Колонка 2: "Для выдачи подотчетным лицам"
    cell_sub1 = table2.rows[1].cells[1]
    para_sub1 = cell_sub1.paragraphs[0]
    para_sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub1 = para_sub1.add_run('Для выдачи\nподотчетным лицам')
    run_sub1.font.size = Pt(8)
    
    # Колонка 3: "Для выдачи заработной платы"
    cell_sub2 = table2.rows[1].cells[2]
    para_sub2 = cell_sub2.paragraphs[0]
    para_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub2 = para_sub2.add_run('Для выдачи\nзаработной платы')
    run_sub2.font.size = Pt(8)
    
    # Строка 3: Значения
    # Колонка 4-5: Сумма
    cell_amount1 = table2.rows[2].cells[3]
    cell_amount1.merge(table2.rows[2].cells[4])
    para_amount1 = cell_amount1.paragraphs[0]
    para_amount1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_amount1 = para_amount1.add_run('{{AMOUNT}}')
    run_amount1.font.size = Pt(8)
    
    # Добавляем пустую строку
    doc.add_paragraph()
    
    # ========== ОСНОВНЫЕ ПОЛЯ ==========
    
    # Выдать
    para_fio = doc.add_paragraph()
    para_fio.add_run('Выдать ').font.size = Pt(9)
    para_fio.add_run('_' * 100).font.size = Pt(9)
    para_fio.add_run(' {{FIO}}').font.size = Pt(9)
    
    # Подпись под ФИО
    para_fio_sub = doc.add_paragraph()
    para_fio_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fio_sub = para_fio_sub.add_run('(фамилия, имя, отчество)')
    run_fio_sub.font.size = Pt(6)
    run_fio_sub.font.name = 'Times New Roman'
    
    # Линия
    para_line1 = doc.add_paragraph()
    para_line1.add_run('_' * 100).font.size = Pt(9)
    
    # Основание
    para_basis = doc.add_paragraph()
    para_basis.add_run('Основание ').font.size = Pt(9)
    para_basis.add_run('_' * 100).font.size = Pt(9)
    para_basis.add_run(' {{BASIS}}').font.size = Pt(9)
    
    # Линия
    para_line2 = doc.add_paragraph()
    para_line2.add_run('_' * 100).font.size = Pt(9)
    
    # Сумма прописью
    para_amount_words = doc.add_paragraph()
    para_amount_words.add_run('Сумма ').font.size = Pt(9)
    para_amount_words.add_run('_' * 100).font.size = Pt(9)
    para_amount_words.add_run(' {{AMOUNT_WORDS}} ').font.size = Pt(9)
    
    # Линия
    para_line3 = doc.add_paragraph()
    para_line3.add_run('_' * 100).font.size = Pt(9)
    
    # Подпись под суммой
    para_amount_sub = doc.add_paragraph()
    para_amount_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_amount_sub = para_amount_sub.add_run('(прописью)')
    run_amount_sub.font.size = Pt(6)
    run_amount_sub.font.name = 'Times New Roman'
    
    # Приложение
    para_attachment = doc.add_paragraph()
    para_attachment.add_run('Приложение').font.size = Pt(9)
    
    # Линия
    para_line4 = doc.add_paragraph()
    para_line4.add_run('_' * 100).font.size = Pt(9)
    
    # Руководитель организации
    para_head = doc.add_paragraph()
    para_head.add_run('Руководитель организации ').font.size = Pt(9)
    para_head.add_run('_' * 50).font.size = Pt(9)
    para_head.add_run(' ИП ').font.size = Pt(11)
    para_head.add_run('_' * 100).font.size = Pt(9)
    para_head.add_run('{{IP_SHORT}}').font.size = Pt(9)
    
    # Линия
    para_line5 = doc.add_paragraph()
    para_line5.add_run('_' * 100).font.size = Pt(9)
    
    # Подпись под руководителем
    para_head_sub = doc.add_paragraph()
    para_head_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_head_sub = para_head_sub.add_run('(должность) ' * 3 + '(подпись) ' * 3 + '(расшифровка подписи)')
    run_head_sub.font.size = Pt(6)
    run_head_sub.font.name = 'Times New Roman'
    
    # Получил
    para_received = doc.add_paragraph()
    para_received.add_run('Получил : ').font.size = Pt(9)
    para_received.add_run('_' * 100).font.size = Pt(9)
    para_received.add_run(' {{AMOUNT_WORDS}} ').font.size = Pt(9)
    
    # Линия
    para_line6 = doc.add_paragraph()
    para_line6.add_run('_' * 100).font.size = Pt(9)
    
    # Подпись под получил
    para_received_sub = doc.add_paragraph()
    para_received_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_received_sub = para_received_sub.add_run('(сумма прописью)')
    run_received_sub.font.size = Pt(6)
    run_received_sub.font.name = 'Times New Roman'
    
    # Дата получения
    para_received_date = doc.add_paragraph()
    para_received_date.add_run('_' * 50).font.size = Pt(9)
    para_received_date.add_run(' {{RECEIVED_DATE}} ').font.size = Pt(9)
    para_received_date.add_run('_' * 100).font.size = Pt(9)
    para_received_date.add_run(' Подпи').font.size = Pt(9)
    
    # Пустая строка
    doc.add_paragraph()
    
    # Паспортные данные
    para_passport = doc.add_paragraph()
    para_passport.add_run('По: ').font.size = Pt(9)
    para_passport.add_run('{{PASSPORT}}').font.size = Pt(9)
    
    # Линия
    para_line7 = doc.add_paragraph()
    para_line7.add_run('_' * 100).font.size = Pt(9)
    
    # Дополнительные паспортные данные
    para_passport2 = doc.add_paragraph()
    run_passport2 = para_passport2.add_run('{{PASSPORT2}}')
    run_passport2.font.size = Pt(9)
    run_passport2.font.underline = True
    
    # Линия
    para_line8 = doc.add_paragraph()
    para_line8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_line8.add_run('_' * 100).font.size = Pt(9)
    
    # Подпись под паспортом
    para_passport_sub = doc.add_paragraph()
    para_passport_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_passport_sub = para_passport_sub.add_run('(серия, номер, кем и когда выдан, адрес постоянного места жительства)')
    run_passport_sub.font.size = Pt(6)
    run_passport_sub.font.name = 'Times New Roman'
    
    # Подпись директора
    para_director = doc.add_paragraph()
    para_director.add_run('_' * 100).font.size = Pt(9)
    para_director.add_run('{{IP_SHORT}}').font.size = Pt(9)
    
    # Кассир
    para_cashier = doc.add_paragraph()
    para_cashier.add_run('Выдал кассир ').font.size = Pt(9)
    para_cashier.add_run('_' * 100).font.size = Pt(9)
    para_cashier.add_run(' ').font.size = Pt(9)
    
    # Подпись под кассиром
    para_cashier_sub = doc.add_paragraph()
    para_cashier_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cashier_sub = para_cashier_sub.add_run('(должность) ' * 2 + '(подпись) ' * 2 + '(расшифровка подписи)')
    run_cashier_sub.font.size = Pt(6)
    run_cashier_sub.font.name = 'Times New Roman'
    
    # Сохраняем документ
    output_path = '/root/.cursor/rko_template_from_html.docx'
    doc.save(output_path)
    print(f'✅ Шаблон РКО создан на основе HTML структуры: {output_path}')
    print('\n📝 Используемые плейсхолдеры:')
    print('  - {{IP}} - Полное имя директора с "ИП"')
    print('  - {{INN}} - ИНН')
    print('  - {{SHOP}} - Адрес магазина')
    print('  - {{DOC_ID}} - Номер документа')
    print('  - {{DATE}} - Дата составления')
    print('  - {{AMOUNT}} - Сумма цифрами')
    print('  - {{FIO}} - ФИО получателя')
    print('  - {{BASIS}} - Основание')
    print('  - {{AMOUNT_WORDS}} - Сумма прописью')
    print('  - {{IP_SHORT}} - Короткое имя директора')
    print('  - {{RECEIVED_DATE}} - Дата получения')
    print('  - {{PASSPORT}} - Паспортные данные')
    print('  - {{PASSPORT2}} - Дополнительные паспортные данные')

if __name__ == '__main__':
    create_rko_from_html_structure()

