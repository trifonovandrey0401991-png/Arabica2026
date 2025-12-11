#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание правильного шаблона РКО КО-2 с фиксированными высотами строк
Это решает проблему съезжающего текста вниз
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_row_height(row, height_pt):
    """Устанавливает фиксированную высоту строки"""
    try:
        tr = row._element
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_pt * 20)))  # Высота в двадцатых долях пункта
        trHeight.set(qn('w:hRule'), 'exact')  # exact = фиксированная высота
        trPr.append(trHeight)
    except Exception as e:
        pass  # Игнорируем ошибки

def set_cell_vertical_alignment(cell, alignment='top'):
    """Устанавливает вертикальное выравнивание в ячейке (top, center, bottom)"""
    try:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), alignment)
        tcPr.append(vAlign)
    except Exception as e:
        print(f"Ошибка установки выравнивания: {e}")

def create_proper_ko2_template():
    """Создает правильный шаблон РКО КО-2 с фиксированными высотами"""
    
    doc = Document()
    
    # Настройка страницы A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    
    # ========== ЗАГОЛОВОК (справа вверху) ==========
    para_header = doc.add_paragraph()
    para_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_header = para_header.add_run('Унифицированная форма № КО-2\nУтверждена постановлением Госкомстата России от 18.08.98 № 88')
    run_header.font.size = Pt(8)
    run_header.font.name = 'Times New Roman'
    
    # ========== ТАБЛИЦА 1: Организация, ИНН, Адрес, Номер документа, Дата ==========
    # Создаем таблицу с фиксированными высотами строк
    table1 = doc.add_table(rows=8, cols=6)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Устанавливаем фиксированные высоты строк (в пунктах)
    row_heights = [15, 10, 15, 10, 10, 10, 10, 20]  # Высоты для каждой строки
    
    for row_idx, row in enumerate(table1.rows):
        if row_idx < len(row_heights):
            set_row_height(row, row_heights[row_idx])
            for cell in row.cells:
                set_cell_vertical_alignment(cell, 'top')  # Выравнивание по верху
    
    # Строка 1: Организация и ИНН (объединенные ячейки 1-2)
    cell_org = table1.rows[0].cells[0]
    cell_org.merge(table1.rows[0].cells[1])
    para_org = cell_org.paragraphs[0]
    para_org.add_run(' {{IP}} ').font.name = 'Arial'
    para_org.add_run('ИНН: ').font.name = 'Arial'
    para_org.add_run('{{INN}}').font.name = 'Arial'
    
    # Линия подчеркивания
    cell_line1 = table1.rows[1].cells[0]
    cell_line1.merge(table1.rows[1].cells[1])
    cell_line1.text = '_______________________________________________________________________'
    
    # Строка 3: Фактический адрес
    cell_addr = table1.rows[2].cells[0]
    cell_addr.merge(table1.rows[2].cells[1])
    para_addr = cell_addr.paragraphs[0]
    run_addr = para_addr.add_run('Фактический адрес: {{SHOP}}')
    run_addr.font.size = Pt(8)
    run_addr.font.underline = True
    
    # Линия подчеркивания для адреса
    cell_line2 = table1.rows[3].cells[0]
    cell_line2.merge(table1.rows[3].cells[1])
    cell_line2.text = '_______________________________________________________________________'
    
    # Строка 4: Подпись "(наименование)"
    cell_name1 = table1.rows[4].cells[0]
    cell_name1.merge(table1.rows[4].cells[1])
    para_name1 = cell_name1.paragraphs[0]
    para_name1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name1 = para_name1.add_run('(наименование)')
    run_name1.font.size = Pt(6)
    
    # Строка 5: Подпись "(наименование организации)"
    cell_name2 = table1.rows[5].cells[0]
    cell_name2.merge(table1.rows[5].cells[1])
    para_name2 = cell_name2.paragraphs[0]
    para_name2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name2 = para_name2.add_run('(наименование организации)')
    run_name2.font.size = Pt(6)
    
    # Строка 6: Пустая
    for col in range(6):
        table1.rows[6].cells[col].text = ''
    
    # Строка 7: Номер документа и Дата
    # Колонка 3: "Номер по порядку" (выравнивание справа)
    cell_num_label = table1.rows[7].cells[2]
    para_num_label = cell_num_label.paragraphs[0]
    para_num_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_num_label = para_num_label.add_run('Номер по порядку')
    run_num_label.font.size = Pt(8)
    
    # Колонка 4-6: Номер документа (объединенные)
    cell_doc_id = table1.rows[7].cells[3]
    cell_doc_id.merge(table1.rows[7].cells[5])
    para_doc_id = cell_doc_id.paragraphs[0]
    para_doc_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_doc_id = para_doc_id.add_run('{{DOC_ID}}')
    run_doc_id.font.size = Pt(8)
    run_doc_id.font.bold = True
    
    # Дата в той же строке, но в отдельном параграфе
    para_date = cell_doc_id.add_paragraph()
    para_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = para_date.add_run('{{DATE}}')
    run_date.font.size = Pt(8)
    run_date.font.bold = True
    
    # Добавляем заголовок "РАСХОДНЫЙ КАССОВЫЙ ОРДЕР"
    para_title = doc.add_paragraph()
    para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = para_title.add_run('РАСХОДНЫЙ КАССОВЫЙ ОРДЕР')
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    run_title.font.name = 'Arial'
    
    # ========== ТАБЛИЦА 2: Сумма (бухгалтерская таблица) ==========
    table2 = doc.add_table(rows=3, cols=6)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Фиксированные высоты для таблицы суммы
    row_heights2 = [15, 10, 20]
    for row_idx, row in enumerate(table2.rows):
        if row_idx < len(row_heights2):
            set_row_height(row, row_heights2[row_idx])
            for cell in row.cells:
                set_cell_vertical_alignment(cell, 'top')
    
    # Заголовки таблицы
    headers = ['', 'Дебет', '', 'Кредит', 'Сумма, руб. коп.', '']
    for col_idx, header in enumerate(headers):
        if col_idx < len(table2.rows[0].cells):
            cell = table2.rows[0].cells[col_idx]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header)
            run.font.size = Pt(8)
            run.font.bold = True
    
    # Подзаголовки
    subheaders = ['', 'код структурного\nподразделения', 'корреспондирующий\nсчет, субсчет', 'код аналитического\nучета', '', '']
    for col_idx, subheader in enumerate(subheaders):
        if col_idx < len(table2.rows[1].cells):
            cell = table2.rows[1].cells[col_idx]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(subheader)
            run.font.size = Pt(7)
    
    # Сумма цифрами в колонке 5
    cell_amount = table2.rows[2].cells[4]
    para_amount = cell_amount.paragraphs[0]
    para_amount.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_amount = para_amount.add_run('{{AMOUNT}}')
    run_amount.font.size = Pt(9)
    run_amount.font.bold = True
    
    # ========== ОСНОВНЫЕ ПОЛЯ (без таблиц, только параграфы) ==========
    doc.add_paragraph()  # Пустая строка
    
    # Выдать
    para_fio = doc.add_paragraph()
    para_fio.add_run('Выдать: ').font.size = Pt(9)
    para_fio.add_run('_' * 50).font.size = Pt(9)
    para_fio.add_run(' {{FIO}}').font.size = Pt(9)
    
    # Основание
    para_basis = doc.add_paragraph()
    para_basis.add_run('Основание: ').font.size = Pt(9)
    para_basis.add_run('_' * 50).font.size = Pt(9)
    para_basis.add_run(' {{BASIS}}').font.size = Pt(9)
    
    # Сумма прописью
    para_amount_words = doc.add_paragraph()
    para_amount_words.add_run('Сумма: ').font.size = Pt(9)
    para_amount_words.add_run('_' * 50).font.size = Pt(9)
    para_amount_words.add_run(' {{AMOUNT_WORDS}}').font.size = Pt(9)
    
    # Приложение
    para_attachment = doc.add_paragraph()
    para_attachment.add_run('Приложение: ').font.size = Pt(9)
    para_attachment.add_run('_' * 50).font.size = Pt(9)
    
    # Руководитель
    para_head = doc.add_paragraph()
    para_head.add_run('Руководитель ').font.size = Pt(9)
    para_head.add_run('_' * 30).font.size = Pt(9)
    para_head.add_run(' {{IP_SHORT}}').font.size = Pt(9)
    
    # Главный бухгалтер
    para_accountant = doc.add_paragraph()
    para_accountant.add_run('Главный бухгалтер ').font.size = Pt(9)
    para_accountant.add_run('_' * 30).font.size = Pt(9)
    
    # Получил
    para_received = doc.add_paragraph()
    para_received.add_run('Получил: ').font.size = Pt(9)
    para_received.add_run('_' * 50).font.size = Pt(9)
    para_received.add_run(' {{AMOUNT_WORDS}}').font.size = Pt(9)
    
    # Дата получения
    para_received_date = doc.add_paragraph()
    para_received_date.add_run('" " ').font.size = Pt(9)
    para_received_date.add_run('{{RECEIVED_DATE}}').font.size = Pt(9)
    para_received_date.add_run(' г. ').font.size = Pt(9)
    para_received_date.add_run('_' * 30).font.size = Pt(9)
    para_received_date.add_run(' Подпись').font.size = Pt(9)
    
    # Паспортные данные
    para_passport = doc.add_paragraph()
    para_passport.add_run('По ').font.size = Pt(9)
    para_passport.add_run('_' * 50).font.size = Pt(9)
    para_passport.add_run(' {{PASSPORT}}').font.size = Pt(9)
    
    # Дополнительные паспортные данные
    para_passport2 = doc.add_paragraph()
    para_passport2.add_run('_' * 100).font.size = Pt(9)
    para_passport2.add_run(' {{PASSPORT2}}').font.size = Pt(9)
    
    # Выдал кассир
    para_cashier = doc.add_paragraph()
    para_cashier.add_run('Выдал кассир ').font.size = Pt(9)
    para_cashier.add_run('_' * 30).font.size = Pt(9)
    para_cashier.add_run(' {{IP_SHORT}}').font.size = Pt(9)
    
    # Сохраняем документ
    output_path = '/root/.cursor/rko_template_ko2_fixed.docx'
    doc.save(output_path)
    print(f'✅ Правильный шаблон КО-2 создан: {output_path}')
    return output_path

if __name__ == '__main__':
    create_proper_ko2_template()

