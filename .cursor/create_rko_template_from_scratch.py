#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание правильного шаблона РКО с нуля, основываясь на понимании структуры документа
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_proper_rko_template():
    """Создает правильный шаблон РКО с правильной структурой"""
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ========== ЗАГОЛОВОК ==========
    title = doc.add_paragraph('РАСХОДНЫЙ КАССОВЫЙ ОРДЕР')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0].font
    title_format.bold = True
    title_format.size = Pt(14)
    
    doc.add_paragraph()  # Пустая строка
    
    # ========== ТАБЛИЦА 1: Организация, ИНН, Адрес, Номер документа, Дата ==========
    table1 = doc.add_table(rows=8, cols=6)
    table1.style = 'Table Grid'
    
    # Настройка ширины колонок
    widths = [Inches(1.5), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.5)]
    for row in table1.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
    
    # Строка 1: Организация и ИНН (объединенные ячейки)
    cell_org = table1.rows[0].cells[0]
    cell_org.merge(table1.rows[0].cells[2])  # Объединяем первые 3 колонки
    para_org = cell_org.paragraphs[0]
    para_org.add_run(' {{IP}} ИНН: {{INN}}')
    
    # Строка 1: Пустые ячейки справа
    for col in range(3, 6):
        table1.rows[0].cells[col].text = ''
    
    # Строка 2: Линия подчеркивания для организации
    cell_line1 = table1.rows[1].cells[0]
    cell_line1.merge(table1.rows[1].cells[2])
    cell_line1.text = '_______________________________________________________________________'
    
    # Строка 3: Фактический адрес
    cell_addr = table1.rows[2].cells[0]
    cell_addr.merge(table1.rows[2].cells[2])
    para_addr = cell_addr.paragraphs[0]
    para_addr.add_run(' Фактический адрес: {{SHOP}}')
    
    # Строка 4: Линия подчеркивания для адреса
    cell_line2 = table1.rows[3].cells[0]
    cell_line2.merge(table1.rows[3].cells[2])
    cell_line2.text = '_______________________________________________________________________'
    
    # Строки 5-7: Пустые строки
    for row_idx in range(4, 7):
        for col_idx in range(6):
            table1.rows[row_idx].cells[col_idx].text = ''
    
    # Строка 8: Номер документа и Дата
    # Номер документа в колонке 2
    cell_doc_id = table1.rows[7].cells[1]
    para_doc_id = cell_doc_id.paragraphs[0]
    para_doc_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_doc_id.add_run('{{DOC_ID}}')
    
    # Дата в колонке 6
    cell_date = table1.rows[7].cells[5]
    para_date = cell_date.paragraphs[0]
    para_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_date.add_run('{{DATE}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # ========== ОСНОВНАЯ ИНФОРМАЦИЯ ==========
    
    # Выдать
    para_fio = doc.add_paragraph()
    para_fio.add_run('  Выдать                                                                {{FIO}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # Основание
    para_basis = doc.add_paragraph()
    para_basis.add_run('  Основание                                                                   {{BASIS}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # Сумма прописью
    para_amount_words = doc.add_paragraph()
    para_amount_words.add_run('  Сумма                                                                      {{AMOUNT_WORDS}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # ========== ТАБЛИЦА 2: Сумма цифрами ==========
    table2 = doc.add_table(rows=3, cols=9)
    table2.style = 'Table Grid'
    
    # Настройка ширины колонок для таблицы суммы
    widths2 = [Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.8), Inches(0.8), Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.5)]
    for row in table2.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths2[idx]
    
    # Строка 1-2: Пустые
    for row_idx in range(2):
        for col_idx in range(9):
            table2.rows[row_idx].cells[col_idx].text = ''
    
    # Строка 3: Сумма в колонках 4 и 5
    cell_amount1 = table2.rows[2].cells[3]
    para_amount1 = cell_amount1.paragraphs[0]
    para_amount1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_amount1.add_run('{{AMOUNT}}')
    
    cell_amount2 = table2.rows[2].cells[4]
    para_amount2 = cell_amount2.paragraphs[0]
    para_amount2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_amount2.add_run('{{AMOUNT}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # ========== ПОДПИСИ ==========
    
    # Руководитель организации
    para_head = doc.add_paragraph()
    para_head.add_run('  Руководитель организации                  ИП                                                       ')
    para_head.add_run('{{IP_SHORT}}')
    
    doc.add_paragraph()  # Пустая строка
    doc.add_paragraph()  # Пустая строка
    
    # Получил
    para_received = doc.add_paragraph()
    para_received.add_run('  Получил :                                                                     {{AMOUNT_WORDS}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # Дата получения
    para_received_date = doc.add_paragraph()
    para_received_date.add_run('                        {{RECEIVED_DATE}}	                                                      Подпи')
    
    doc.add_paragraph()  # Пустая строка
    
    # Паспортные данные
    para_passport = doc.add_paragraph()
    para_passport.add_run('  По: {{PASSPORT}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # Дополнительные паспортные данные
    para_passport2 = doc.add_paragraph()
    para_passport2.add_run('{{PASSPORT2}}')
    
    doc.add_paragraph()  # Пустая строка
    
    # Подпись директора (если нужно)
    para_director = doc.add_paragraph()
    para_director.add_run('                                                                                                     ')
    para_director.add_run('{{IP_SHORT}}')
    
    # Сохраняем документ
    output_path = '/root/.cursor/rko_template_proper.docx'
    doc.save(output_path)
    print(f'✅ Правильный шаблон РКО создан: {output_path}')
    print('\n📝 Используемые плейсхолдеры:')
    print('  - {{IP}} - Полное имя директора с "ИП"')
    print('  - {{INN}} - ИНН')
    print('  - {{SHOP}} - Адрес магазина')
    print('  - {{DOC_ID}} - Номер документа')
    print('  - {{DATE}} - Дата составления')
    print('  - {{FIO}} - ФИО получателя')
    print('  - {{BASIS}} - Основание')
    print('  - {{AMOUNT_WORDS}} - Сумма прописью')
    print('  - {{AMOUNT}} - Сумма цифрами')
    print('  - {{IP_SHORT}} - Короткое имя директора')
    print('  - {{RECEIVED_DATE}} - Дата получения')
    print('  - {{PASSPORT}} - Паспортные данные')
    print('  - {{PASSPORT2}} - Дополнительные паспортные данные')
    print('\n💡 Шаблон создан с правильной структурой. Откройте в Word для финальной настройки форматирования.')

if __name__ == '__main__':
    create_proper_rko_template()

